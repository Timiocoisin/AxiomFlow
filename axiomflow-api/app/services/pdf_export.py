from __future__ import annotations

import io
import logging
import os
import platform
import tempfile
from datetime import datetime
from pathlib import Path

import fitz  # PyMuPDF

from ..core.config import settings
from ..core.config_manager import config_manager
from .font_subset import (
    optimize_font_embedding,
    collect_used_characters,
    create_font_subset,
)

logger = logging.getLogger(__name__)


# 可选：DOCX 导出依赖（python-docx）
try:  # pragma: no cover - 依赖可能不存在
    from docx import Document  # type: ignore[import]
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import]
    from docx.shared import Pt  # type: ignore[import]

    _DOCX_AVAILABLE = True
except Exception:  # noqa: BLE001
    Document = None  # type: ignore[assignment]
    WD_ALIGN_PARAGRAPH = None  # type: ignore[assignment]
    Pt = None  # type: ignore[assignment]
    _DOCX_AVAILABLE = False


def _get_cjk_font_path(*, subset_chars: set[str] | None = None) -> str | None:
    """
    获取 CJK 字体路径（优先级）：
    1. 用户通过配置/环境变量 AXIOMFLOW_PDF_FONT 指定的字体
    2. 系统自动查找的常用 CJK 字体
    
    Args:
        subset_chars: 如果提供，则创建字体子集（仅包含这些字符）
    
    Returns:
        字体文件路径（可能是子集化后的临时文件）
    """
    # 优先级 1：用户指定字体
    # 首选 app/core/config.py 中的 Settings（支持 .env: AXIOMFLOW_PDF_FONT 或 PDF_FONT）
    raw_path = settings.pdf_font or ""
    if not raw_path:
        # 兼容老的 ConfigManager 配置（以及通过环境变量覆盖的 AXIOMFLOW_PDF_FONT）
        raw_path = str(config_manager.get("AXIOMFLOW_PDF_FONT", "") or "")
    # 允许在 env 中使用相对路径（如 "./fonts/msyh.ttc" 或 "fonts/msyh.ttc"）：
    # 这里统一把相对路径转换为以项目根目录为基准的绝对路径，
    # 避免 Celery / API 进程的工作目录不一致导致找不到字体。
    if raw_path:
        font_path = Path(raw_path)
        if not font_path.is_absolute():
            repo_root = Path(__file__).resolve().parents[2]
            font_path = (repo_root / raw_path).resolve()
        path = str(font_path)
    else:
        path = ""

    if path:
        logger.info(f"CJK font path resolved to: {path}")
    else:
        logger.info("AXIOMFLOW_PDF_FONT not set or empty; will fall back to system fonts")

    if path and Path(path).exists():
        # 如果需要子集化且提供了字符集合，创建子集
        if subset_chars:
            try:
                subset_path = create_font_subset(path, subset_chars)
                if subset_path != path:
                    return subset_path
            except Exception:
                # 如果子集化失败，返回原字体
                pass
        logger.info(f"Using CJK font from AXIOMFLOW_PDF_FONT: {path}")
        return path

    # 优先级 2：系统自动查找常用 CJK 字体
    system = platform.system()
    common_cjk_fonts = []

    if system == "Windows":
        # Windows 系统字体路径
        font_dirs = [
            Path(os.getenv("WINDIR", "C:/Windows")) / "Fonts",
        ]
        common_cjk_fonts = [
            "msyh.ttc",  # 微软雅黑
            "simsun.ttc",  # 宋体
            "simhei.ttf",  # 黑体
            "simkai.ttf",  # 楷体
            "simli.ttf",  # 隶书
        ]
    elif system == "Darwin":  # macOS
        font_dirs = [
            Path("/System/Library/Fonts"),
            Path("/Library/Fonts"),
            Path.home() / "Library/Fonts",
        ]
        common_cjk_fonts = [
            "STHeiti Light.ttc",  # 黑体
            "STSong.ttc",  # 宋体
            "STKaiti.ttc",  # 楷体
            "PingFang.ttc",  # 苹方
        ]
    else:  # Linux
        font_dirs = [
            Path("/usr/share/fonts"),
            Path("/usr/local/share/fonts"),
            Path.home() / ".fonts",
            Path.home() / ".local/share/fonts",
        ]
        common_cjk_fonts = [
            "NotoSansCJK-Regular.ttc",
            "NotoSerifCJK-Regular.ttc",
            "SourceHanSansCN-Regular.otf",
            "SourceHanSerifCN-Regular.otf",
        ]

    # 尝试查找常用 CJK 字体
    for font_dir in font_dirs:
        if not font_dir.exists():
            continue
        for font_name in common_cjk_fonts:
            font_path = font_dir / font_name
            if font_path.exists():
                found_path = str(font_path)
                # 如果需要子集化且提供了字符集合，创建子集
                if subset_chars:
                    try:
                        subset_path = create_font_subset(found_path, subset_chars)
                        if subset_path != found_path:
                            return subset_path
                    except Exception:
                        # 如果子集化失败，返回原字体
                        pass
                return found_path

    # 通用查找：查找包含 CJK 相关关键字的字体文件
    for font_dir in font_dirs:
        if not font_dir.exists():
            continue
        for ext in ["*.ttf", "*.ttc", "*.otf"]:
            for font_file in font_dir.glob(ext):
                name_lower = font_file.name.lower()
                # 匹配常见的中文字体标识
                if any(keyword in name_lower for keyword in ["cjk", "han", "chinese", "noto", "source"]):
                    found_path = str(font_file)
                    # 如果需要子集化且提供了字符集合，创建子集
                    if subset_chars:
                        try:
                            subset_path = create_font_subset(found_path, subset_chars)
                            if subset_path != found_path:
                                return subset_path
                        except Exception:
                            # 如果子集化失败，返回原字体
                            pass
                    return found_path

    return None


def _insert_block(
    page: fitz.Page,
    block: dict,
    *,
    use_translation: bool,
    font_path: str | None,
    cover_original: bool,
) -> None:
    """
    在 PDF 页面上插入文本块（译文或原文）
    
    优化：
    - 支持从 block 中提取原始字体大小
    - 自适应字体大小处理文本溢出
    - 智能文本对齐方式推断
    - 文本编码验证和清理
    """
    if block.get("is_header_footer"):
        return
    if block.get("is_footnote"):
        return
    bbox = block.get("bbox") or {}
    if not bbox:
        return
    rect = fitz.Rect(float(bbox["x0"]), float(bbox["y0"]), float(bbox["x1"]), float(bbox["y1"]))

    src = (block.get("text") or "").strip()
    dst = (block.get("translation") or "").strip()
    
    # 如果使用译文但译文为空，回退到原文
    if use_translation:
        if dst:
            text = dst
        elif src:
            # 译文为空但有原文，使用原文并记录警告
            logger.warning(
                f"译文为空，使用原文: block_id={block.get('id', 'unknown')}, "
                f"type={block.get('type', 'unknown')}"
            )
            text = src
        else:
            # 原文和译文都为空，跳过
            logger.debug(f"文本块为空，跳过: block_id={block.get('id', 'unknown')}")
            return
    else:
        text = src
        if not text:
            return
    
    # 文本编码验证和清理：检测并修复常见的编码问题
    try:
        # 确保文本是有效的UTF-8字符串
        if isinstance(text, bytes):
            text = text.decode('utf-8', errors='replace')
        elif not isinstance(text, str):
            text = str(text)
        
        # 检测并过滤掉明显的乱码字符（连续的替换字符或问号）
        # 如果文本中超过30%是替换字符或问号，可能是编码问题
        if len(text) > 0:
            replacement_chars = text.count('\ufffd')  # Unicode替换字符
            question_marks = text.count('?')
            total_suspicious = replacement_chars + question_marks
            if total_suspicious / len(text) > 0.3:
                logger.warning(
                    f"检测到可能的编码问题: block_id={block.get('id', 'unknown')}, "
                    f"可疑字符比例={total_suspicious/len(text):.2%}, "
                    f"原文长度={len(src)}, 译文长度={len(dst)}"
                )
                # 如果原文存在且译文看起来有问题，尝试使用原文
                if use_translation and dst and src and len(src) > 0:
                    logger.info(f"使用原文替代可能有问题的译文: block_id={block.get('id', 'unknown')}")
                    text = src
    except Exception as e:
        logger.warning(f"文本编码验证失败: {e}, block_id={block.get('id', 'unknown')}")
        # 如果验证失败，继续使用原文本

    t = (block.get("type") or "paragraph").lower()
    # 对 figure/table 块不做文本覆盖（保留原图/原表）
    if t in ("figure", "table"):
        return

    # 优先使用 block 中保存的原始字体大小，否则根据类型推断
    fontsize = block.get("font_size")
    if not fontsize:
        if t == "heading":
            fontsize = 13
        elif t == "caption":
            fontsize = 9
        elif t == "formula":
            fontsize = 11
        else:
            fontsize = 10

    # 高保真导出：在原页上覆盖原文字区域（白底），再叠加译文
    # 注意：即使译文为空但使用了原文，也要覆盖原区域
    if cover_original and use_translation:
        page.draw_rect(rect, color=None, fill=(1, 1, 1), overlay=True)

    # 推断文本对齐方式（基于文本在矩形中的位置）
    align = fitz.TEXT_ALIGN_LEFT  # 默认左对齐
    # 如果 block 中有保存的对齐方式，使用它
    if "align" in block:
        align_map = {
            "left": fitz.TEXT_ALIGN_LEFT,
            "center": fitz.TEXT_ALIGN_CENTER,
            "right": fitz.TEXT_ALIGN_RIGHT,
            "justify": fitz.TEXT_ALIGN_JUSTIFY,
        }
        align = align_map.get(block["align"], fitz.TEXT_ALIGN_LEFT)

    # 字体嵌入：优先使用自定义字体文件，确保 CJK 字符正确显示
    try:
        # 验证字体文件是否存在（如果提供了字体路径）
        if font_path and not Path(font_path).exists():
            logger.warning(f"字体文件不存在: {font_path}，将使用系统字体")
            font_path = None
        
        if font_path:
            # 使用自定义字体文件，PyMuPDF 会自动嵌入
            # 尝试插入文本，如果溢出则缩小字体
            overflow = page.insert_textbox(
                rect,
                text,
                fontfile=font_path,
                fontsize=fontsize,
                color=(0, 0, 0),
                align=align,
                render_mode=0,  # 填充模式
            )
            # 如果文本溢出，尝试缩小字体
            if overflow > 0 and fontsize > 6:
                # 计算缩放因子（基于溢出量）
                scale = max(0.7, 1.0 - (overflow / 100))
                new_fontsize = max(6, fontsize * scale)
                page.insert_textbox(
                    rect,
                    text,
                    fontfile=font_path,
                    fontsize=new_fontsize,
                    color=(0, 0, 0),
                    align=align,
                    render_mode=0,
                )
        else:
            # 兜底：使用内置字体（可能不支持 CJK，但至少不会出错）
            # 对于包含CJK字符的文本，尝试使用支持CJK的系统字体
            has_cjk = any(
                0x4E00 <= ord(char) <= 0x9FFF or  # 中文
                0x3040 <= ord(char) <= 0x309F or  # 日文平假名
                0x30A0 <= ord(char) <= 0x30FF or  # 日文片假名
                0xAC00 <= ord(char) <= 0xD7AF      # 韩文
                for char in text
            )
            
            if has_cjk:
                # 尝试使用支持CJK的字体名称
                cjk_fonts = ["china-s", "china-ss", "japan", "korea"]
                font_inserted = False
                for cjk_font in cjk_fonts:
                    try:
                        overflow = page.insert_textbox(
                            rect,
                            text,
                            fontname=cjk_font,
                            fontsize=fontsize,
                            color=(0, 0, 0),
                            align=align,
                            render_mode=0,
                        )
                        font_inserted = True
                        if overflow > 0 and fontsize > 6:
                            scale = max(0.7, 1.0 - (overflow / 100))
                            new_fontsize = max(6, fontsize * scale)
                            page.insert_textbox(
                                rect,
                                text,
                                fontname=cjk_font,
                                fontsize=new_fontsize,
                                color=(0, 0, 0),
                                align=align,
                                render_mode=0,
                            )
                        break
                    except Exception:
                        continue
                
                if not font_inserted:
                    # 如果所有CJK字体都失败，回退到helv
                    logger.warning(f"无法使用CJK字体插入文本，回退到helv: block_id={block.get('id', 'unknown')}")
                    overflow = page.insert_textbox(
                        rect,
                        text,
                        fontname="helv",
                        fontsize=fontsize,
                        color=(0, 0, 0),
                        align=align,
                        render_mode=0,
                    )
                    if overflow > 0 and fontsize > 6:
                        scale = max(0.7, 1.0 - (overflow / 100))
                        new_fontsize = max(6, fontsize * scale)
                        page.insert_textbox(
                            rect,
                            text,
                            fontname="helv",
                            fontsize=new_fontsize,
                            color=(0, 0, 0),
                            align=align,
                            render_mode=0,
                        )
            else:
                # 非CJK文本，使用helv字体
                overflow = page.insert_textbox(
                    rect,
                    text,
                    fontname="helv",
                    fontsize=fontsize,
                    color=(0, 0, 0),
                    align=align,
                    render_mode=0,
                )
                # 如果文本溢出，尝试缩小字体
                if overflow > 0 and fontsize > 6:
                    scale = max(0.7, 1.0 - (overflow / 100))
                    new_fontsize = max(6, fontsize * scale)
                    page.insert_textbox(
                        rect,
                        text,
                        fontname="helv",
                        fontsize=new_fontsize,
                        color=(0, 0, 0),
                        align=align,
                        render_mode=0,
                    )
    except Exception as e:
        # 如果插入失败，记录错误并尝试多种回退方案
        logger.warning(
            f"文本插入失败: {e}, block_id={block.get('id', 'unknown')}, "
            f"text_preview={text[:50] if len(text) > 50 else text}, "
            f"font_path={font_path}, has_cjk={any(0x4E00 <= ord(c) <= 0x9FFF for c in text)}"
        )
        
        # 回退方案1：尝试使用更小的字体
        try:
            fallback_size = max(6, fontsize * 0.8)
            if font_path and Path(font_path).exists():
                page.insert_textbox(rect, text, fontfile=font_path, fontsize=fallback_size, color=(0, 0, 0), align=align)
                logger.info(f"使用回退方案1成功: block_id={block.get('id', 'unknown')}")
                return
        except Exception:
            pass
        
        # 回退方案2：如果字体文件失败，尝试使用系统CJK字体
        try:
            has_cjk = any(
                0x4E00 <= ord(char) <= 0x9FFF or
                0x3040 <= ord(char) <= 0x309F or
                0x30A0 <= ord(char) <= 0x30FF or
                0xAC00 <= ord(char) <= 0xD7AF
                for char in text
            )
            
            if has_cjk:
                # 尝试使用系统CJK字体
                cjk_fonts = ["china-s", "china-ss", "japan", "korea"]
                for cjk_font in cjk_fonts:
                    try:
                        page.insert_textbox(rect, text, fontname=cjk_font, fontsize=fallback_size, color=(0, 0, 0), align=align)
                        logger.info(f"使用回退方案2成功 (CJK字体 {cjk_font}): block_id={block.get('id', 'unknown')}")
                        return
                    except Exception:
                        continue
            
            # 回退方案3：使用helv字体（可能不支持CJK，但至少能显示英文）
            try:
                page.insert_textbox(rect, text, fontname="helv", fontsize=fallback_size, color=(0, 0, 0), align=align)
                logger.warning(f"使用回退方案3 (helv字体，CJK字符可能无法显示): block_id={block.get('id', 'unknown')}")
                return
            except Exception:
                pass
        except Exception:
            pass
        
        # 所有回退方案都失败，记录严重错误
        fallback_size = max(6, fontsize * 0.8)  # 确保变量已定义
        logger.error(
            f"文本插入所有方案均失败，跳过该块: block_id={block.get('id', 'unknown')}, "
            f"text_length={len(text)}, bbox={bbox}, "
            f"原文={src[:50] if src else 'None'}, 译文={dst[:50] if dst else 'None'}"
        )
        
        # 如果原文存在且译文插入失败，尝试插入原文作为最后手段
        if use_translation and src and src != text:
            try:
                logger.info(f"尝试插入原文作为最后手段: block_id={block.get('id', 'unknown')}")
                page.insert_textbox(rect, src, fontname="helv", fontsize=fallback_size, color=(0, 0, 0), align=align)
                return
            except Exception:
                pass
        
        # 如果所有方法都失败，至少保留白底覆盖，避免显示原文
        # 这样用户知道这里应该有内容但插入失败了


def _convert_to_pdfa(
    pdf_bytes: bytes,
    *,
    part: int = 2,
    conformance: str = "B",
) -> bytes:
    """
    将 PDF 转换为 PDF/A 格式（PDF/A-2B 兼容性）
    
    使用 pikepdf 设置 PDF/A 元数据和 OutputIntent，提高长期兼容性。
    """
    try:
        import pikepdf
    except ImportError:  # pragma: no cover - 依赖可选
        # 如果 pikepdf 未安装，返回原 PDF（不强制要求）
        return pdf_bytes

    # 使用临时文件进行转换
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_in:
        tmp_in.write(pdf_bytes)
        tmp_in_path = tmp_in.name

    try:
        # 使用 pikepdf 打开并转换
        pdf = pikepdf.Pdf.open(tmp_in_path)

        # 添加 PDF/A 元数据
        if not pdf.docinfo or pdf.docinfo.get("/Title") is None:
            pdf.docinfo["/Title"] = "Translated Document"

        # 设置 PDF/A 兼容性元数据（默认为 PDF/A-2B，可通过参数细粒度控制）
        with pdf.open_metadata() as meta:
            # 保留原有元数据
            if pdf.docinfo:
                try:
                    meta.load_from_docinfo(pdf.docinfo)
                except Exception:
                    pass
            
            # 设置 PDF/A 标识
            meta["pdfaid:part"] = str(part)
            meta["pdfaid:conformance"] = conformance.upper()

        # 创建 OutputIntent（颜色配置文件）
        output_intent = pikepdf.Dictionary(
            {
                pikepdf.Name("/Type"): pikepdf.Name("/OutputIntent"),
                pikepdf.Name("/S"): pikepdf.Name("/GTS_PDFA1"),
                pikepdf.Name("/OutputConditionIdentifier"): "sRGB IEC61966-2.1",
                pikepdf.Name("/RegistryName"): "http://www.color.org",
                pikepdf.Name("/Info"): "sRGB IEC61966-2.1",
            }
        )

        # 添加到 PDF 根对象
        output_intents_name = pikepdf.Name("/OutputIntents")
        if output_intents_name not in pdf.Root:
            pdf.Root[output_intents_name] = pikepdf.Array([output_intent])
        else:
            pdf.Root[output_intents_name].append(output_intent)

        # 保存为 PDF/A
        tmp_out_path = tmp_in_path + ".pdfa"
        pdf.save(tmp_out_path, linearize=True)
        pdf.close()

        # 读取转换后的 PDF
        with open(tmp_out_path, "rb") as f:
            result = f.read()

        # 清理临时文件
        os.unlink(tmp_in_path)
        os.unlink(tmp_out_path)

        return result

    except Exception:
        # 如果转换失败，返回原 PDF
        if os.path.exists(tmp_in_path):
            os.unlink(tmp_in_path)
        return pdf_bytes


def build_translated_pdf(
    structured: dict,
    *,
    bilingual: bool,
    subset_fonts: bool = True,
    convert_to_pdfa: bool = False,
    pdfa_part: int | None = None,
    pdfa_conformance: str = "B",
) -> bytes:
    """
    构建翻译后的 PDF（高保真版本）
    
    Args:
        structured: 结构化文档 JSON
        bilingual: 是否为双语模式（原文+译文交替页）
        subset_fonts: 是否进行字体子集化（减小文件大小，默认启用）
        convert_to_pdfa: 是否转换为 PDF/A 格式（提高兼容性，默认禁用）
    
    Returns:
        PDF 文件的字节数据
    """
    # 收集使用的字符（用于字体子集化）
    used_chars = None
    if subset_fonts:
        try:
            used_chars = collect_used_characters(structured)
        except Exception:
            # 如果收集失败，继续使用完整字体
            used_chars = None
    
    # 获取字体路径（如果需要，会创建子集）
    font_path = _get_cjk_font_path(subset_chars=used_chars)
    pages = structured.get("pages", [])
    source_pdf_path = structured.get("document", {}).get("source_pdf_path")
    if not source_pdf_path:
        raise RuntimeError("Missing document.source_pdf_path for high-fidelity PDF export")

    src = fitz.open(source_pdf_path)
    page_count = src.page_count
    
    # 跟踪是否创建了临时子集字体（需要最后清理）
    temp_subset_font = None
    if font_path:
        # 检查是否是临时文件（子集化创建的）
        temp_subset_font_path = Path(font_path)
        # 检查路径是否在临时目录中
        if (temp_subset_font_path.parent.name.startswith("tmp") or 
            "temp" in str(temp_subset_font_path.parent).lower() or
            temp_subset_font_path.parent == Path(tempfile.gettempdir())):
            temp_subset_font = font_path

    # 先复制原文页，保证图片/表格线条、书签、注释等保留
    base = fitz.open()
    base.insert_pdf(src, start_at=0, to_page=page_count - 1)

    if bilingual:
        # 双语：再复制一份原文页作为译文页底板
        base.insert_pdf(src, start_at=0, to_page=page_count - 1)

    # 使用增强的PDF结构保留模块
    try:
        from .pdf_structure_preserve import preserve_pdf_structure

        preserve_pdf_structure(src, base, bilingual=bilingual)
    except ImportError:
        # 回退到基础实现
        logger.warning("PDF结构保留增强模块未找到，使用基础实现")
        
        # 保留原 PDF 的元数据（标题、作者等）
        try:
            if src.metadata:
                base.set_metadata(src.metadata)
        except Exception:
            pass

        # 保留原 PDF 的书签（目录）
        try:
            if src.is_pdf:
                toc = src.get_toc(simple=False)
                if toc:
                    base.set_toc(toc)
        except Exception:
            pass

        # 保留原 PDF 的批注和表单字段（基础实现）
        try:
            for src_page_idx in range(page_count):
                src_page = src.load_page(src_page_idx)
                base_page = base.load_page(src_page_idx)
                
                try:
                    base_page.copy_annotations(src_page)
                except Exception:
                    pass
                
                if bilingual:
                    base_dual_page = base.load_page(src_page_idx + page_count)
                    try:
                        base_dual_page.copy_annotations(src_page)
                    except Exception:
                        pass
        except Exception as e:
            logger.debug(f"复制批注/表单字段时出错（已忽略）: {e}")

    # 链接通常由 insert_pdf 自动复制，无需额外处理

    # 在"译文页底板"上覆盖并叠加译文
    total_blocks = 0
    blocks_with_translation = 0
    blocks_without_translation = 0
    blocks_with_corrupted_translation = 0
    
    for p in pages:
        page_index = int(p.get("index") or 0)
        blocks = list(p.get("blocks", []))
        if page_index < 0 or page_index >= page_count:
            continue

        if bilingual:
            target_page = base.load_page(page_index + page_count)
        else:
            target_page = base.load_page(page_index)

        for b in blocks:
            total_blocks += 1
            src_text = (b.get("text") or "").strip()
            dst_text = (b.get("translation") or "").strip()
            
            # 统计翻译情况
            if dst_text:
                blocks_with_translation += 1
                # 检查译文是否包含大量问号（可能是乱码）
                if dst_text.count('?') / len(dst_text) > 0.2 if len(dst_text) > 0 else False:
                    blocks_with_corrupted_translation += 1
                    logger.warning(
                        f"检测到可能损坏的译文: page={page_index + 1}, "
                        f"block_id={b.get('id', 'unknown')}, "
                        f"type={b.get('type', 'unknown')}, "
                        f"原文长度={len(src_text)}, 译文长度={len(dst_text)}, "
                        f"问号比例={dst_text.count('?')/len(dst_text):.2%}, "
                        f"原文预览={src_text[:50] if len(src_text) > 50 else src_text}, "
                        f"译文预览={dst_text[:50] if len(dst_text) > 50 else dst_text}"
                    )
            elif src_text:
                blocks_without_translation += 1
                logger.debug(
                    f"译文为空，将使用原文: page={page_index + 1}, "
                    f"block_id={b.get('id', 'unknown')}, "
                    f"type={b.get('type', 'unknown')}"
                )
            
            _insert_block(
                target_page,
                b,
                use_translation=True,
                font_path=font_path,
                cover_original=True,
            )
    
    # 记录统计信息
    logger.info(
        f"PDF导出统计: 总块数={total_blocks}, "
        f"有译文={blocks_with_translation}, "
        f"无译文={blocks_without_translation}, "
        f"损坏译文={blocks_with_corrupted_translation}, "
        f"字体路径={font_path}"
    )
    
    if blocks_with_corrupted_translation > 0:
        logger.warning(
            f"检测到 {blocks_with_corrupted_translation} 个可能损坏的译文块，"
            f"建议检查翻译服务是否正常工作"
        )

    # 输出：mono 直接返回；dual 需要交替页序
    if not bilingual:
        # 优化字体嵌入（子集化）
        optimize_font_embedding(base, subset_fonts=subset_fonts, used_chars=used_chars)
        
        pdf_bytes = base.tobytes(deflate=True, garbage=4)
        base.close()
        src.close()
        
        # 清理临时子集字体文件
        if temp_subset_font and Path(temp_subset_font).exists():
            try:
                Path(temp_subset_font).unlink()
            except Exception:
                pass

        # PDF/A 转换（可选，可配置 part/conformance）
        if convert_to_pdfa:
            pdf_bytes = _convert_to_pdfa(
                pdf_bytes,
                part=pdfa_part or 2,
                conformance=pdfa_conformance,
            )

        return pdf_bytes

    # 双语模式：重新排列页序（原文页 + 译文页交替）
    out = fitz.open()
    for i in range(page_count):
        out.insert_pdf(base, from_page=i, to_page=i)
        out.insert_pdf(base, from_page=i + page_count, to_page=i + page_count)
    
    # 优化字体嵌入（子集化）
    optimize_font_embedding(out, subset_fonts=subset_fonts, used_chars=used_chars)
    
    pdf_bytes = out.tobytes(deflate=True, garbage=4)
    out.close()
    base.close()
    src.close()
    
    # 清理临时子集字体文件
    if temp_subset_font and Path(temp_subset_font).exists():
        try:
            Path(temp_subset_font).unlink()
        except Exception:
            pass

    # PDF/A 转换（可选，可配置 part/conformance）
    if convert_to_pdfa:
        pdf_bytes = _convert_to_pdfa(
            pdf_bytes,
            part=pdfa_part or 2,
            conformance=pdfa_conformance,
        )

    return pdf_bytes


def export_pdf(
    structured: dict,
    *,
    kind: str,
    subset_fonts: bool = True,
    convert_to_pdfa: bool = False,
    pdfa_part: int | None = None,
    pdfa_conformance: str = "B",
) -> tuple[str, bytes]:
    """
    导出 PDF 文件
    
    Args:
        structured: 结构化文档 JSON
        kind: 'mono' | 'dual'
        subset_fonts: 是否进行字体子集化（默认启用）
        convert_to_pdfa: 是否转换为 PDF/A 格式（默认禁用）
        pdfa_part: PDF/A part（1/2/3），默认 2
        pdfa_conformance: PDF/A conformance（"A" | "B" | "U"），默认 "B"
    
    Returns:
        (filename, pdf_bytes) 元组
    """
    now = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    doc_id = structured.get("document", {}).get("id", "document")
    
    if kind == "dual":
        filename = f"{doc_id}_{now}_dual.pdf"
        data = build_translated_pdf(
            structured,
            bilingual=True,
            subset_fonts=subset_fonts,
            convert_to_pdfa=convert_to_pdfa,
            pdfa_part=pdfa_part,
            pdfa_conformance=pdfa_conformance,
        )
        return filename, data
    
    filename = f"{doc_id}_{now}_mono.pdf"
    data = build_translated_pdf(
        structured,
        bilingual=False,
        subset_fonts=subset_fonts,
        convert_to_pdfa=convert_to_pdfa,
        pdfa_part=pdfa_part,
        pdfa_conformance=pdfa_conformance,
    )
    return filename, data


def build_translated_docx(
    structured: dict,
    *,
    bilingual: bool = False,
) -> bytes:
    """
    构建 DOCX 文档（使用 python-docx，可选依赖）。

    - 根据 block.type 选择合适的段落/标题样式；
    - 支持 bilingual 模式（原文 + 译文分段展示）。
    """
    if not _DOCX_AVAILABLE or Document is None:  # type: ignore[truthy-function]
        raise RuntimeError(
            "DOCX 导出需要安装 python-docx，请运行: pip install python-docx"
        )

    from io import BytesIO

    doc = Document()  # type: ignore[call-arg]

    pages = structured.get("pages", [])
    # 按 reading_order 全局排序，避免页内乱序
    all_blocks: list[dict] = []
    for p in pages:
        for b in p.get("blocks", []):
            all_blocks.append(b)
    all_blocks.sort(key=lambda b: int(b.get("reading_order", 0)))

    fig_no = 0
    tbl_no = 0

    for blk in all_blocks:
        t = (blk.get("type") or "paragraph").lower()
        src = (blk.get("text") or "").strip()
        dst = (blk.get("translation") or "").strip()

        if blk.get("is_header_footer") or blk.get("is_footnote"):
            continue

        val = dst or src
        if not val:
            continue

        # 标题：使用 heading 样式，可根据 section_level 调整级别
        if t == "heading":
            level = int(blk.get("section_level") or 1)
            level = max(1, min(level, 3))
            para = doc.add_heading(val, level=level)  # type: ignore[call-arg]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # type: ignore[assignment]
            continue

        # 图表标题
        if t == "caption":
            if bilingual and dst:
                para = doc.add_paragraph()  # type: ignore[call-arg]
                run = para.add_run(src)
                run.italic = True
                para2 = doc.add_paragraph()  # type: ignore[call-arg]
                run2 = para2.add_run(dst)
                run2.italic = True
            else:
                para = doc.add_paragraph()  # type: ignore[call-arg]
                run = para.add_run(val)
                run.italic = True
            continue

        # 图/表占位
        if t in ("figure", "table"):
            if t == "figure":
                fig_no += 1
                label = f"Figure {fig_no}"
            else:
                tbl_no += 1
                label = f"Table {tbl_no}"
            para = doc.add_paragraph()  # type: ignore[call-arg]
            run = para.add_run(f"[{label}]")
            run.bold = True
            continue

        # 公式：使用等宽或预格式风格（简单处理）
        if t == "formula":
            para = doc.add_paragraph()  # type: ignore[call-arg]
            run = para.add_run(val)
            run.font.name = "Consolas"  # type: ignore[assignment]
            run.font.size = Pt(10)  # type: ignore[call-arg]
            continue

        # 普通段落
        if bilingual and dst:
            # 原文
            para_src = doc.add_paragraph()  # type: ignore[call-arg]
            para_src.add_run(src)
            # 译文
            para_dst = doc.add_paragraph()  # type: ignore[call-arg]
            run_dst = para_dst.add_run(dst)
            run_dst.font.size = Pt(11)  # type: ignore[call-arg]
        else:
            para = doc.add_paragraph()  # type: ignore[call-arg]
            run = para.add_run(val)
            run.font.size = Pt(11)  # type: ignore[call-arg]

    bio = BytesIO()
    doc.save(bio)  # type: ignore[call-arg]
    return bio.getvalue()


def export_docx(
    structured: dict,
    *,
    bilingual: bool = False,
) -> tuple[str, bytes]:
    """
    导出 DOCX 文件。

    Args:
        structured: 结构化文档 JSON
        bilingual: 是否导出原文+译文双语（逐段展示）
    """
    now = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    doc_id = structured.get("document", {}).get("id", "document")
    filename = f"{doc_id}_{now}.docx"
    data = build_translated_docx(structured, bilingual=bilingual)
    return filename, data

